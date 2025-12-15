"""
Writing module - Generates methods and results sections using Claude Opus.
"""

import os
import json
import tempfile
from typing import Dict, List, Tuple, Optional
import pandas as pd
import requests
from prompts import PROMPTS
from journal_formats import JOURNAL_FORMATS

class ResultsWriter:
    """Generates publication-ready methods and results sections."""
    
    def __init__(self, openrouter_api_key: str):
        self.api_key = openrouter_api_key
        self.base_url = "https://openrouter.ai/api/v1/chat/completions"
        self.model_id = "anthropic/claude-opus-4-5"
        self.input_cost = 5.00  # per 1M tokens
        self.output_cost = 25.00  # per 1M tokens
    
    def _call_opus(self, messages: List[Dict], temperature: float = 0.3) -> Tuple[str, float]:
        """Call Claude Opus via OpenRouter."""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://stats-council.streamlit.app",
            "X-Title": "Stats Council"
        }
        
        payload = {
            "model": self.model_id,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": 8192
        }
        
        try:
            response = requests.post(self.base_url, headers=headers, json=payload, timeout=180)
            response.raise_for_status()
            result = response.json()
            
            content = result['choices'][0]['message']['content']
            
            # Calculate cost
            usage = result.get('usage', {})
            input_tokens = usage.get('prompt_tokens', 0)
            output_tokens = usage.get('completion_tokens', 0)
            cost = (input_tokens * self.input_cost / 1_000_000) + (output_tokens * self.output_cost / 1_000_000)
            
            return content, cost
            
        except Exception as e:
            return f"Error: {str(e)}", 0.0
    
    def generate_results_document(self, df: pd.DataFrame, analysis_plan: str,
                                  execution_results: str, figures: List[bytes],
                                  tables: Dict[str, pd.DataFrame], review: str,
                                  journal: str, study_design: str) -> Tuple[str, float]:
        """
        Generate complete methods and results sections.
        
        Returns:
            Tuple of (document_path, cost)
        """
        # Get journal-specific formatting
        journal_format = JOURNAL_FORMATS.get(journal, JOURNAL_FORMATS['Generic'])
        
        # Determine reporting guideline
        reporting_guideline = self._get_reporting_guideline(study_design)
        
        # Prepare table summaries for context
        table_summaries = {}
        for name, table in tables.items():
            table_summaries[name] = table.to_string()[:2000]  # Limit for context
        
        # Generate methods section
        methods_prompt = PROMPTS['methods_writing'].format(
            analysis_plan=analysis_plan,
            study_design=study_design,
            reporting_guideline=reporting_guideline,
            journal_format=json.dumps(journal_format, indent=2),
            sample_size=len(df)
        )
        
        messages = [
            {"role": "system", "content": PROMPTS['writing_system']},
            {"role": "user", "content": methods_prompt}
        ]
        
        methods_text, methods_cost = self._call_opus(messages)
        
        # Generate results section
        results_prompt = PROMPTS['results_writing'].format(
            execution_results=execution_results,
            table_summaries=json.dumps(table_summaries, indent=2),
            num_figures=len(figures),
            journal_format=json.dumps(journal_format, indent=2),
            reporting_guideline=reporting_guideline
        )
        
        messages = [
            {"role": "system", "content": PROMPTS['writing_system']},
            {"role": "user", "content": results_prompt}
        ]
        
        results_text, results_cost = self._call_opus(messages)
        
        # Generate figure legends
        legends_prompt = PROMPTS['figure_legends'].format(
            num_figures=len(figures),
            execution_results=execution_results[:3000],
            journal_format=json.dumps(journal_format, indent=2)
        )
        
        messages = [
            {"role": "system", "content": PROMPTS['writing_system']},
            {"role": "user", "content": legends_prompt}
        ]
        
        legends_text, legends_cost = self._call_opus(messages)
        
        # Generate limitations paragraph
        limitations_prompt = PROMPTS['limitations_writing'].format(
            study_design=study_design,
            review=review,
            analysis_plan=analysis_plan
        )
        
        messages = [
            {"role": "system", "content": PROMPTS['writing_system']},
            {"role": "user", "content": limitations_prompt}
        ]
        
        limitations_text, limitations_cost = self._call_opus(messages)
        
        total_cost = methods_cost + results_cost + legends_cost + limitations_cost
        
        # Create Word document
        doc_path = self._create_word_document(
            methods_text, results_text, legends_text, limitations_text,
            tables, journal_format
        )
        
        return doc_path, total_cost
    
    def _get_reporting_guideline(self, study_design: str) -> str:
        """Get appropriate reporting guideline for study design."""
        guidelines = {
            'Retrospective Cohort': 'STROBE (Strengthening the Reporting of Observational Studies in Epidemiology)',
            'Prospective Cohort': 'STROBE',
            'Case-Control': 'STROBE',
            'Cross-sectional': 'STROBE',
            'RCT': 'CONSORT (Consolidated Standards of Reporting Trials)',
            'Case Series': 'CARE (Case Report Guidelines)',
            'Prediction Model': 'TRIPOD (Transparent Reporting of a Multivariable Prediction Model)',
            'Auto-detect': 'STROBE'
        }
        return guidelines.get(study_design, 'STROBE')
    
    def _create_word_document(self, methods: str, results: str, legends: str,
                             limitations: str, tables: Dict[str, pd.DataFrame],
                             journal_format: dict) -> str:
        """Create Word document using python-docx library."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Output path
        output_path = tempfile.mktemp(suffix='.docx')

        try:
            # Create document
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)

            # Title
            title = doc.add_heading('Statistical Analysis Results', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Methods section
            doc.add_heading('Methods', level=1)
            self._add_formatted_text(doc, methods)

            # Results section
            doc.add_heading('Results', level=1)
            self._add_formatted_text(doc, results)

            # Add tables if any
            if tables:
                doc.add_page_break()
                doc.add_heading('Tables', level=1)
                for table_name, table_df in tables.items():
                    doc.add_heading(table_name, level=2)
                    self._add_dataframe_table(doc, table_df)
                    doc.add_paragraph()

            # Figure legends
            doc.add_heading('Figure Legends', level=1)
            self._add_formatted_text(doc, legends)

            # Limitations
            doc.add_heading('Limitations', level=1)
            self._add_formatted_text(doc, limitations)

            # Save document
            doc.save(output_path)
            return output_path

        except Exception as e:
            # Fallback: create simple text file
            fallback_path = tempfile.mktemp(suffix='.txt')
            with open(fallback_path, 'w') as f:
                f.write("METHODS\n\n")
                f.write(methods)
                f.write("\n\nRESULTS\n\n")
                f.write(results)
                f.write("\n\nFIGURE LEGENDS\n\n")
                f.write(legends)
                f.write("\n\nLIMITATIONS\n\n")
                f.write(limitations)

            return fallback_path

    def _add_formatted_text(self, doc, text: str):
        """Add formatted text to document, handling markdown-style formatting."""
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check for headers
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            else:
                # Regular paragraph - handle bold text marked with **
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:  # Odd indices are between ** markers
                            run.bold = True

    def _add_dataframe_table(self, doc, df: pd.DataFrame):
        """Add a pandas DataFrame as a table to the document."""
        # Create table with header row
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'

        # Add header row
        header_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            header_cells[i].text = str(col_name)

        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
